import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
import os
import time
import shutil
import io
from pathlib import Path
from PIL import Image, ImageDraw, ImageTk

# ── API KEY FIJA ───────────────────────────────────────────────────────────────
GEMINI_API_KEY = " PUT HERE YOUR API KEY "
# ──────────────────────────────────────────────────────────────────────────────

DOWNLOADS = str(Path.home() / "Downloads")

C = {
    "fondo":      "#060F1E",
    "panel":      "#0B1A2E",
    "card":       "#0F2240",
    "borde":      "#163260",
    "borde2":     "#1E4A8A",
    "acento":     "#2563EB",
    "acento2":    "#38BDF8",
    "acento3":    "#0EA5E9",
    "btn":        "#2563EB",
    "btn_hover":  "#1D4ED8",
    "texto":      "#EFF6FF",
    "texto2":     "#7BA4CC",
    "texto3":     "#4A7BA8",
    "exito":      "#10B981",
    "error":      "#EF4444",
    "warn":       "#F59E0B",
    "tag_bg":     "#0F2A50",
}

EXTENSIONES = {
    ".jpg": "imagen", ".jpeg": "imagen", ".png": "imagen",
    ".tiff": "imagen", ".tif": "imagen", ".bmp": "imagen", ".webp": "imagen",
    ".pdf": "pdf",
    ".xlsx": "excel", ".xls": "excel",
    ".docx": "word", ".doc": "word",
}

FORMATOS_DESTINO = {
    "imagen": ["Excel", "PDF", "Word", "Imagen"],
    "pdf":    ["Excel", "Imágenes", "Word", "PDF"],
    "excel":  ["PDF", "Word", "Imágenes", "Excel"],
    "word":   ["PDF", "Excel", "Imágenes", "Word"],
}

ICONOS_TIPO = {
    "imagen": "🖼️",
    "pdf":    "📄",
    "excel":  "📊",
    "word":   "📝",
}


def crear_icono_app():
    """Ícono: documento con flecha de conversión."""
    size = 256
    img = Image.new("RGBA", (size, size), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)

    # Fondo circular azul degradado
    for i in range(size // 2):
        t = i / (size // 2)
        r = int(6 + (37 - 6) * t)
        g = int(15 + (99 - 15) * t)
        b = int(30 + (235 - 30) * t)
        alpha = 255
        draw.ellipse([i, i, size - i, size - i], fill=(r, g, b, alpha))

    # Documento izquierdo (blanco)
    doc_color = (220, 235, 255)
    fold_color = (150, 190, 255)
    # Cuerpo
    draw.rounded_rectangle([30, 55, 105, 185], radius=6, fill=doc_color)
    # Doblez esquina
    draw.polygon([(85, 55), (105, 55), (105, 78)], fill=fold_color)
    draw.polygon([(85, 55), (85, 78), (105, 78)], fill=(180, 210, 255))
    # Líneas de texto
    for y in [100, 115, 130, 145, 160]:
        ancho = 45 if y == 160 else 55
        draw.rounded_rectangle([42, y, 42 + ancho, y + 6], radius=2, fill=(150, 190, 240))

    # Flecha central (→)
    cx = 128
    cy = 125
    # Cuerpo flecha
    draw.rounded_rectangle([cx - 18, cy - 6, cx + 10, cy + 6], radius=4, fill=(56, 189, 248))
    # Punta flecha
    draw.polygon([
        (cx + 8, cy - 14),
        (cx + 26, cy),
        (cx + 8, cy + 14),
    ], fill=(56, 189, 248))

    # Documento derecho (destino, con acento celeste)
    draw.rounded_rectangle([150, 55, 225, 185], radius=6, fill=(200, 230, 255))
    draw.polygon([(205, 55), (225, 55), (225, 78)], fill=(100, 170, 240))
    draw.polygon([(205, 55), (205, 78), (225, 78)], fill=(140, 200, 255))
    # Líneas
    for y in [100, 115, 130, 145, 160]:
        ancho = 45 if y == 160 else 55
        draw.rounded_rectangle([162, y, 162 + ancho, y + 6], radius=2, fill=(56, 189, 248))

    return img


def tipo_archivo(ruta):
    return EXTENSIONES.get(Path(ruta).suffix.lower(), None)


# ── Botón redondeado moderno ──────────────────────────────────────────────────
class BtnModerno(tk.Canvas):
    def __init__(self, parent, texto, cmd, ancho=180, alto=44, primario=True, **kwargs):
        super().__init__(parent, width=ancho, height=alto,
                         bg=C["fondo"], highlightthickness=0, **kwargs)
        self.cmd = cmd
        self.ancho = ancho
        self.alto = alto
        self.texto = texto
        self.primario = primario
        self._activo = True
        self._draw(hover=False)
        self.bind("<Enter>", lambda e: self._draw(hover=True))
        self.bind("<Leave>", lambda e: self._draw(hover=False))
        self.bind("<Button-1>", self._click)

    def _draw(self, hover=False):
        self.delete("all")
        r = 10
        w, h = self.ancho, self.alto
        if not self._activo:
            bg = "#1A2A40"
            fg = C["texto3"]
        elif self.primario:
            bg = C["btn_hover"] if hover else C["btn"]
            fg = C["texto"]
        else:
            bg = C["card"] if hover else C["panel"]
            fg = C["acento2"]

        # Rounded rect
        self.create_arc(0, 0, 2*r, 2*r, start=90, extent=90, fill=bg, outline="")
        self.create_arc(w-2*r, 0, w, 2*r, start=0, extent=90, fill=bg, outline="")
        self.create_arc(0, h-2*r, 2*r, h, start=180, extent=90, fill=bg, outline="")
        self.create_arc(w-2*r, h-2*r, w, h, start=270, extent=90, fill=bg, outline="")
        self.create_rectangle(r, 0, w-r, h, fill=bg, outline="")
        self.create_rectangle(0, r, w, h-r, fill=bg, outline="")
        self.create_text(w//2, h//2, text=self.texto,
                         fill=fg, font=("Segoe UI", 10, "bold"))

    def _click(self, e):
        if self._activo and self.cmd:
            self.cmd()

    def set_activo(self, val):
        self._activo = val
        self._draw()


# ── Tarjeta de archivo ────────────────────────────────────────────────────────
class TarjetaArchivo(tk.Frame):
    def __init__(self, parent, ruta, on_remove, **kwargs):
        super().__init__(parent, bg=C["tag_bg"], padx=8, pady=4, **kwargs)
        self.ruta = ruta
        tipo = tipo_archivo(ruta)
        icono = ICONOS_TIPO.get(tipo, "📁")
        nombre = Path(ruta).name
        nombre_corto = nombre[:30] + "..." if len(nombre) > 30 else nombre

        tk.Label(self, text=icono, bg=C["tag_bg"], fg=C["texto"],
                 font=("Segoe UI", 11)).pack(side="left", padx=(0, 4))
        tk.Label(self, text=nombre_corto, bg=C["tag_bg"], fg=C["texto"],
                 font=("Segoe UI", 9)).pack(side="left")
        tk.Label(self, text="✕", bg=C["tag_bg"], fg=C["error"],
                 font=("Segoe UI", 9, "bold"), cursor="hand2").pack(side="left", padx=(8, 0))
        self.winfo_children()[-1].bind("<Button-1>", lambda e: on_remove(self))


# ── APP PRINCIPAL ─────────────────────────────────────────────────────────────
class ConvertidorIA(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Convertidor IA")
        self.geometry("620x720")
        self.resizable(False, False)
        self.configure(bg=C["fondo"])
        self._archivos = []
        self._conversion = False
        self._set_icono()
        self._build()

    def _set_icono(self):
        try:
            ic = crear_icono_app()
            self._ic_tk = ImageTk.PhotoImage(ic.resize((32, 32)))
            self.iconphoto(True, self._ic_tk)
        except Exception:
            pass

    # ── UI ────────────────────────────────────────────────────────────────────
    def _build(self):
        # Header
        hdr = tk.Frame(self, bg=C["panel"], height=72)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)

        try:
            ic_hdr = crear_icono_app().resize((44, 44))
            self._ic_hdr = ImageTk.PhotoImage(ic_hdr)
            tk.Label(hdr, image=self._ic_hdr, bg=C["panel"]).pack(side="left", padx=18)
        except Exception:
            pass

        ft = tk.Frame(hdr, bg=C["panel"])
        ft.pack(side="left", pady=10)
        tk.Label(ft, text="CONVERTIDOR IA", font=("Segoe UI", 17, "bold"),
                 fg=C["acento2"], bg=C["panel"]).pack(anchor="w")
        tk.Label(ft, text="Conversión inteligente con IA · Preciso y rápido",
                 font=("Segoe UI", 8), fg=C["texto2"], bg=C["panel"]).pack(anchor="w")

        tk.Frame(self, bg=C["borde"], height=1).pack(fill="x")

        # Contenido
        cont = tk.Frame(self, bg=C["fondo"])
        cont.pack(fill="both", expand=True, padx=28, pady=18)

        # ── Zona drop ──
        tk.Label(cont, text="ARCHIVOS DE ENTRADA", font=("Segoe UI", 8, "bold"),
                 fg=C["texto2"], bg=C["fondo"]).pack(anchor="w", pady=(0, 6))

        self.zona = tk.Canvas(cont, height=130, bg=C["card"],
                               highlightthickness=1, highlightbackground=C["borde2"],
                               cursor="hand2")
        self.zona.pack(fill="x")
        self.zona.bind("<Button-1>", lambda e: self._elegir_archivos())
        self.zona.bind("<Enter>", lambda e: self._zona_highlight(True))
        self.zona.bind("<Leave>", lambda e: self._zona_highlight(False))

        # Drag & Drop
        try:
            self.zona.drop_target_register("DND_Files")
            self.zona.dnd_bind("<<Drop>>", self._on_drop)
        except Exception:
            pass  # tkdnd opcional

        self._dibujar_zona()

        # Lista de archivos seleccionados
        self.frame_archivos = tk.Frame(cont, bg=C["fondo"])
        self.frame_archivos.pack(fill="x", pady=(8, 0))

        # Botón elegir archivos
        frame_btn_elegir = tk.Frame(cont, bg=C["fondo"])
        frame_btn_elegir.pack(fill="x", pady=(6, 0))
        BtnModerno(frame_btn_elegir, "📂  Elegir archivos", self._elegir_archivos,
                   ancho=160, alto=36, primario=False).pack(side="left")
        BtnModerno(frame_btn_elegir, "🗑  Limpiar todo", self._limpiar_todo,
                   ancho=140, alto=36, primario=False).pack(side="left", padx=(8, 0))

        tk.Frame(cont, bg=C["borde"], height=1).pack(fill="x", pady=14)

        # ── Formato destino ──
        tk.Label(cont, text="CONVERTIR A", font=("Segoe UI", 8, "bold"),
                 fg=C["texto2"], bg=C["fondo"]).pack(anchor="w", pady=(0, 8))

        self.frame_fmt = tk.Frame(cont, bg=C["fondo"])
        self.frame_fmt.pack(fill="x")

        self.fmt_var = tk.StringVar()
        self.lbl_fmt_hint = tk.Label(self.frame_fmt,
                                      text="← Primero seleccioná archivos",
                                      font=("Segoe UI", 9, "italic"),
                                      fg=C["texto3"], bg=C["fondo"])
        self.lbl_fmt_hint.pack(anchor="w")

        tk.Frame(cont, bg=C["borde"], height=1).pack(fill="x", pady=14)

        # ── Opciones extra ──
        tk.Label(cont, text="OPCIONES", font=("Segoe UI", 8, "bold"),
                 fg=C["texto2"], bg=C["fondo"]).pack(anchor="w", pady=(0, 8))

        frame_opts = tk.Frame(cont, bg=C["fondo"])
        frame_opts.pack(fill="x")

        self.opt_carpeta = tk.BooleanVar(value=False)
        tk.Checkbutton(frame_opts, text="Guardar en carpeta separada por conversión",
                        variable=self.opt_carpeta,
                        fg=C["texto2"], bg=C["fondo"], selectcolor=C["card"],
                        activebackground=C["fondo"], activeforeground=C["texto"],
                        font=("Segoe UI", 9)).pack(anchor="w")

        tk.Frame(cont, bg=C["borde"], height=1).pack(fill="x", pady=14)

        # ── Progreso ──
        self.lbl_estado = tk.Label(cont, text="", font=("Segoe UI", 9),
                                    fg=C["texto2"], bg=C["fondo"])
        self.lbl_estado.pack(anchor="w")

        style = ttk.Style()
        style.theme_use("default")
        style.configure("IA.Horizontal.TProgressbar",
                         troughcolor=C["card"], background=C["acento2"],
                         darkcolor=C["acento2"], lightcolor=C["acento2"],
                         bordercolor=C["borde"], thickness=6)

        self.barra = ttk.Progressbar(cont, style="IA.Horizontal.TProgressbar",
                                      mode="indeterminate", length=564)

        # ── Botón principal ──
        self.btn_main = BtnModerno(cont, "▶   COMENZAR CONVERSIÓN",
                                    self._iniciar, ancho=564, alto=52)
        self.btn_main.pack(pady=(16, 0))

        # Footer
        tk.Frame(self, bg=C["borde"], height=1).pack(fill="x")
        tk.Label(self, text="Convertidor IA  •  Impulsado por Google Gemini  •  v1.0",
                 font=("Segoe UI", 7), fg=C["texto3"], bg=C["panel"]).pack(pady=6)

    def _dibujar_zona(self, highlight=False):
        self.zona.delete("all")
        w = self.zona.winfo_reqwidth() or 564
        h = 130
        color = C["acento"] if highlight else C["borde2"]

        # Borde punteado
        seg = 8
        for x in range(0, w, seg * 2):
            self.zona.create_line(x, 1, min(x+seg, w), 1, fill=color, width=1)
            self.zona.create_line(x, h-1, min(x+seg, w), h-1, fill=color, width=1)
        for y in range(0, h, seg * 2):
            self.zona.create_line(1, y, 1, min(y+seg, h), fill=color, width=1)
            self.zona.create_line(w-1, y, w-1, min(y+seg, h), fill=color, width=1)

        icono_color = C["acento2"] if highlight else C["texto3"]
        self.zona.create_text(w//2, 48, text="⬆", font=("Segoe UI", 22), fill=icono_color)
        self.zona.create_text(w//2, 82, text="Arrastrá archivos aquí o hacé click para seleccionar",
                               fill=C["texto2"] if not highlight else C["acento2"],
                               font=("Segoe UI", 9))
        self.zona.create_text(w//2, 102, text="📄 PDF  🖼️ Imágenes  📊 Excel  📝 Word",
                               fill=C["texto3"], font=("Segoe UI", 8))

    def _zona_highlight(self, on):
        self._dibujar_zona(highlight=on)

    def _elegir_archivos(self):
        tipos = [
            ("Todos los soportados", "*.pdf *.jpg *.jpeg *.png *.tiff *.tif *.bmp *.webp *.xlsx *.xls *.docx *.doc"),
            ("PDF", "*.pdf"),
            ("Imágenes", "*.jpg *.jpeg *.png *.tiff *.tif *.bmp *.webp"),
            ("Excel", "*.xlsx *.xls"),
            ("Word", "*.docx *.doc"),
        ]
        rutas = filedialog.askopenfilenames(filetypes=tipos)
        for r in rutas:
            self._agregar_archivo(r)

    def _on_drop(self, event):
        rutas = self.tk.splitlist(event.data)
        for r in rutas:
            self._agregar_archivo(r)

    def _agregar_archivo(self, ruta):
        if ruta in [a["ruta"] for a in self._archivos]:
            return
        tipo = tipo_archivo(ruta)
        if not tipo:
            messagebox.showwarning("Formato no soportado", f"No se puede convertir:\n{Path(ruta).name}")
            return
        self._archivos.append({"ruta": ruta, "tipo": tipo})
        self._actualizar_lista()
        self._actualizar_formatos()

    def _actualizar_lista(self):
        for w in self.frame_archivos.winfo_children():
            w.destroy()

        for info in self._archivos:
            tipo = info["tipo"]
            icono = ICONOS_TIPO.get(tipo, "📁")
            nombre = Path(info["ruta"]).name
            nombre_corto = nombre[:38] + "..." if len(nombre) > 38 else nombre

            fila = tk.Frame(self.frame_archivos, bg=C["tag_bg"])
            fila.pack(fill="x", pady=2)

            tk.Label(fila, text=icono, bg=C["tag_bg"], fg=C["texto"],
                     font=("Segoe UI", 10), width=2).pack(side="left", padx=(6, 4))
            tk.Label(fila, text=nombre_corto, bg=C["tag_bg"], fg=C["texto"],
                     font=("Segoe UI", 9), anchor="w").pack(side="left", fill="x", expand=True)

            ruta = info["ruta"]
            btn_x = tk.Label(fila, text=" ✕ ", bg=C["tag_bg"], fg=C["error"],
                              font=("Segoe UI", 9, "bold"), cursor="hand2")
            btn_x.pack(side="right", padx=4)
            btn_x.bind("<Button-1>", lambda e, r=ruta: self._quitar_archivo(r))

    def _quitar_archivo(self, ruta):
        self._archivos = [a for a in self._archivos if a["ruta"] != ruta]
        self._actualizar_lista()
        self._actualizar_formatos()

    def _limpiar_todo(self):
        self._archivos = []
        self._actualizar_lista()
        self._actualizar_formatos()

    def _actualizar_formatos(self):
        for w in self.frame_fmt.winfo_children():
            w.destroy()

        if not self._archivos:
            tk.Label(self.frame_fmt, text="← Primero seleccioná archivos",
                     font=("Segoe UI", 9, "italic"), fg=C["texto3"],
                     bg=C["fondo"]).pack(anchor="w")
            return

        # Calcular formatos comunes a todos los archivos seleccionados
        sets = [set(FORMATOS_DESTINO.get(a["tipo"], [])) for a in self._archivos]
        comunes = sets[0].intersection(*sets[1:]) if len(sets) > 1 else sets[0]

        if not comunes:
            tk.Label(self.frame_fmt,
                     text="⚠️  Los archivos seleccionados no tienen formatos de destino en común",
                     font=("Segoe UI", 9), fg=C["warn"], bg=C["fondo"]).pack(anchor="w")
            return

        # Mantener orden original
        orden = ["Excel", "PDF", "Word", "Imágenes", "Imagen"]
        opciones = [o for o in orden if o in comunes]

        self.fmt_var.set(opciones[0])

        for op in opciones:
            rb = tk.Radiobutton(self.frame_fmt, text=op, variable=self.fmt_var, value=op,
                                 font=("Segoe UI", 10),
                                 fg=C["texto"], bg=C["card"],
                                 selectcolor=C["acento"],
                                 activebackground=C["card"],
                                 activeforeground=C["acento2"],
                                 relief="flat", padx=16, pady=8,
                                 indicatoron=False, borderwidth=0)
            rb.pack(side="left", padx=(0, 10))

    # ── Conversión ────────────────────────────────────────────────────────────
    def _iniciar(self):
        if not self._archivos:
            messagebox.showwarning("Sin archivos", "Primero seleccioná archivos para convertir.")
            return
        if not self.fmt_var.get():
            messagebox.showwarning("Sin formato", "Seleccioná el formato de destino.")
            return
        if self._conversion:
            return
        threading.Thread(target=self._run_conversion, daemon=True).start()

    def _run_conversion(self):
        self._conversion = True
        self.btn_main.set_activo(False)
        self.barra.pack(fill="x", pady=(4, 0))
        self.barra.start(10)

        destino = self.fmt_var.get().lower()
        total = len(self._archivos)
        errores = []

        for idx, info in enumerate(self._archivos, 1):
            ruta = info["ruta"]
            tipo = info["tipo"]
            nombre = Path(ruta).name
            self._estado(f"[{idx}/{total}] Convirtiendo: {nombre[:45]}...")

            try:
                salida = self._convertir_uno(ruta, tipo, destino)
                self._estado(f"✅ [{idx}/{total}] {nombre[:40]} → listo", C["exito"])
            except Exception as e:
                errores.append((nombre, str(e)))
                self._estado(f"❌ [{idx}/{total}] Error en {nombre[:35]}", C["error"])

        self.barra.stop()
        self.barra.pack_forget()

        if not errores:
            self._estado(f"✅ {total} archivo(s) convertido(s) con éxito → Carpeta Descargas", C["exito"])
            messagebox.showinfo("¡Listo!", f"Se convirtieron {total} archivo(s).\nGuardados en: {DOWNLOADS}")
        else:
            ok = total - len(errores)
            self._estado(f"⚠️  {ok}/{total} convertidos. {len(errores)} con error.", C["warn"])
            detalles = "\n".join([f"• {n}: {e[:60]}" for n, e in errores])
            messagebox.showwarning("Conversión parcial",
                                    f"{ok} de {total} archivos convertidos.\n\nErrores:\n{detalles}")

        self._conversion = False
        self.btn_main.set_activo(True)

    def _estado(self, texto, color=None):
        self.lbl_estado.config(text=texto, fg=color or C["texto2"])

    def _nombre_salida(self, origen, nueva_ext):
        base = Path(origen).stem
        ruta = os.path.join(DOWNLOADS, f"{base}{nueva_ext}")
        n = 1
        while os.path.exists(ruta):
            ruta = os.path.join(DOWNLOADS, f"{base}_{n}{nueva_ext}")
            n += 1
        return ruta

    def _convertir_uno(self, origen, tipo, destino):
        if tipo == "imagen" and destino == "excel":
            return self._img_excel(origen)
        elif tipo == "imagen" and destino in ("pdf",):
            return self._img_pdf(origen)
        elif tipo == "imagen" and destino == "word":
            return self._img_word(origen)
        elif tipo == "imagen" and destino in ("imagen", "imágenes"):
            return self._img_img(origen)
        elif tipo == "pdf" and destino == "excel":
            return self._pdf_excel(origen)
        elif tipo == "pdf" and destino in ("imágenes", "imagen"):
            return self._pdf_imgs(origen)
        elif tipo == "pdf" and destino == "word":
            return self._pdf_word(origen)
        elif tipo == "pdf" and destino == "pdf":
            salida = self._nombre_salida(origen, "_copia.pdf")
            shutil.copy2(origen, salida)
            return salida
        elif tipo == "excel" and destino == "pdf":
            return self._excel_pdf(origen)
        elif tipo == "excel" and destino in ("imágenes", "imagen"):
            return self._excel_imgs(origen)
        elif tipo == "excel" and destino == "word":
            return self._excel_word(origen)
        elif tipo == "word" and destino == "pdf":
            return self._word_pdf(origen)
        elif tipo == "word" and destino in ("imágenes", "imagen"):
            return self._word_imgs(origen)
        elif tipo == "word" and destino == "excel":
            return self._word_excel(origen)
        else:
            raise Exception(f"Conversión {tipo} → {destino} no soportada")

    # ── Gemini OCR ────────────────────────────────────────────────────────────
    def _gemini_ocr(self, imagen_pil):
        from google import genai
        cliente = genai.Client(api_key=GEMINI_API_KEY)
        prompt = """Analizá esta imagen que contiene una tabla escaneada.
Extraé TODOS los datos respetando exactamente la estructura de columnas y filas.
Devolvé SOLO el CSV puro:
- Primera fila: nombres de columnas
- Separador: punto y coma (;)
- Sin explicaciones ni bloques de código
- Incluí todas las filas sin omitir ninguna"""
        for intento in range(5):
            try:
                resp = cliente.models.generate_content(
                    model="gemini-2.0-flash", contents=[prompt, imagen_pil])
                texto = resp.text
                if "```" in texto:
                    texto = "\n".join([l for l in texto.split("\n") if not l.startswith("```")])
                return texto.strip()
            except Exception as e:
                if intento < 4:
                    time.sleep(15)
                else:
                    raise e

    def _csv_df(self, texto):
        import pandas as pd
        lineas = [l for l in texto.strip().split("\n") if l.strip()]
        if not lineas:
            return None
        filas = [l.split(";") for l in lineas]
        mc = max(len(f) for f in filas)
        filas = [f + [""] * (mc - len(f)) for f in filas]
        return pd.DataFrame(filas[1:], columns=filas[0])

    # ── CONVERSIONES ──────────────────────────────────────────────────────────

    def _img_excel(self, origen):
        img = Image.open(origen)
        csv_txt = self._gemini_ocr(img)
        df = self._csv_df(csv_txt)
        if df is None:
            raise Exception("No se detectó tabla en la imagen")
        salida = self._nombre_salida(origen, ".xlsx")
        df.to_excel(salida, index=False)
        return salida

    def _img_pdf(self, origen):
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate
        from reportlab.platypus import Image as RLImg
        img = Image.open(origen)
        w, h = img.size
        orient = landscape(A4) if w > h else A4
        salida = self._nombre_salida(origen, ".pdf")
        doc = SimpleDocTemplate(salida, pagesize=orient,
                                 leftMargin=20, rightMargin=20, topMargin=20, bottomMargin=20)
        rl = RLImg(origen, width=orient[0]-40, height=orient[1]-40, kind="proportional")
        doc.build([rl])
        return salida

    def _img_word(self, origen):
        from docx import Document
        from docx.shared import Inches
        doc = Document()
        doc.add_picture(origen, width=Inches(6))
        salida = self._nombre_salida(origen, ".docx")
        doc.save(salida)
        return salida

    def _img_img(self, origen):
        """Convierte imagen a PNG optimizado."""
        img = Image.open(origen).convert("RGB")
        salida = self._nombre_salida(origen, "_convertida.png")
        img.save(salida, "PNG", optimize=True)
        return salida

    def _pdf_excel(self, origen):
        import fitz
        base = Path(origen).stem
        carpeta = os.path.join(DOWNLOADS, f"{base}_excel")
        os.makedirs(carpeta, exist_ok=True)
        doc = fitz.open(origen)
        total = len(doc)
        for i, pag in enumerate(doc, 1):
            self._estado(f"⏳ PDF → Excel: página {i}/{total}...")
            pix = pag.get_pixmap(matrix=fitz.Matrix(2, 2))
            img = Image.open(io.BytesIO(pix.tobytes("jpeg")))
            csv_txt = self._gemini_ocr(img)
            df = self._csv_df(csv_txt)
            if df is not None:
                df.to_excel(os.path.join(carpeta, f"pagina_{i:03d}.xlsx"), index=False)
            time.sleep(2)
        doc.close()
        return carpeta

    def _pdf_imgs(self, origen):
        import fitz
        base = Path(origen).stem
        carpeta = os.path.join(DOWNLOADS, f"{base}_imagenes")
        os.makedirs(carpeta, exist_ok=True)
        doc = fitz.open(origen)
        total = len(doc)
        for i, pag in enumerate(doc, 1):
            self._estado(f"⏳ PDF → Imagen: página {i}/{total}...")
            pix = pag.get_pixmap(matrix=fitz.Matrix(2, 2))
            pix.save(os.path.join(carpeta, f"pagina_{i:03d}.png"))
        doc.close()
        return carpeta

    def _pdf_word(self, origen):
        import fitz
        from docx import Document
        from docx.shared import Inches
        doc_pdf = fitz.open(origen)
        doc_w = Document()
        total = len(doc_pdf)
        for i, pag in enumerate(doc_pdf, 1):
            self._estado(f"⏳ PDF → Word: página {i}/{total}...")
            pix = pag.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_io = io.BytesIO(pix.tobytes("png"))
            doc_w.add_picture(img_io, width=Inches(6))
            if i < total:
                doc_w.add_page_break()
        doc_pdf.close()
        salida = self._nombre_salida(origen, ".docx")
        doc_w.save(salida)
        return salida

    def _excel_pdf(self, origen):
        import openpyxl
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors
        from reportlab.lib.units import cm

        wb = openpyxl.load_workbook(origen, data_only=True)
        salida = self._nombre_salida(origen, ".pdf")
        sz = landscape(A4)
        doc = SimpleDocTemplate(salida, pagesize=sz,
                                 leftMargin=cm, rightMargin=cm,
                                 topMargin=1.5*cm, bottomMargin=1.5*cm)
        els = []
        sty = getSampleStyleSheet()

        for idx, hoja in enumerate(wb.sheetnames):
            ws = wb[hoja]
            self._estado(f"⏳ Excel → PDF: hoja {hoja}")
            els.append(Paragraph(f"<b>{hoja}</b>", sty["Heading2"]))
            datos = [[str(c) if c is not None else "" for c in r]
                     for r in ws.iter_rows(values_only=True)
                     if any(c is not None for c in r)]
            if datos:
                nc = len(datos[0])
                aw = (sz[0] - 2*cm) / nc
                t = Table(datos, colWidths=[aw]*nc, repeatRows=1)
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#2563EB")),
                    ("TEXTCOLOR", (0,0), (-1,0), colors.white),
                    ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
                    ("FONTSIZE", (0,0), (-1,-1), 7),
                    ("GRID", (0,0), (-1,-1), 0.3, colors.HexColor("#163260")),
                    ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#EBF4FF")]),
                    ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
                    ("TOPPADDING", (0,0), (-1,-1), 3),
                    ("BOTTOMPADDING", (0,0), (-1,-1), 3),
                ]))
                els.append(t)
            if idx < len(wb.sheetnames) - 1:
                els.append(PageBreak())

        doc.build(els)
        return salida

    def _excel_imgs(self, origen):
        import fitz
        pdf_temp = os.path.join(DOWNLOADS, "_temp_excelpdf.pdf")
        # Reusa _excel_pdf pero con nombre temp
        orig_fn = self._nombre_salida
        self._nombre_salida = lambda o, e: pdf_temp
        try:
            self._excel_pdf(origen)
        finally:
            self._nombre_salida = orig_fn

        base = Path(origen).stem
        carpeta = os.path.join(DOWNLOADS, f"{base}_imagenes")
        os.makedirs(carpeta, exist_ok=True)
        doc = fitz.open(pdf_temp)
        total = len(doc)
        for i, pag in enumerate(doc, 1):
            self._estado(f"⏳ Excel → Imagen: página {i}/{total}...")
            pix = pag.get_pixmap(matrix=fitz.Matrix(2, 2))
            pix.save(os.path.join(carpeta, f"pagina_{i:03d}.png"))
        doc.close()
        if os.path.exists(pdf_temp):
            os.remove(pdf_temp)
        return carpeta

    def _excel_word(self, origen):
        import openpyxl
        from docx import Document
        wb = openpyxl.load_workbook(origen, data_only=True)
        doc = Document()
        for hoja in wb.sheetnames:
            self._estado(f"⏳ Excel → Word: hoja {hoja}")
            ws = wb[hoja]
            doc.add_heading(hoja, level=1)
            datos = [[str(c) if c is not None else "" for c in r]
                     for r in ws.iter_rows(values_only=True)
                     if any(c is not None for c in r)]
            if datos:
                t = doc.add_table(rows=len(datos), cols=len(datos[0]))
                t.style = "Table Grid"
                for i, fila in enumerate(datos):
                    for j, cel in enumerate(fila):
                        t.cell(i, j).text = cel
            doc.add_page_break()
        salida = self._nombre_salida(origen, ".docx")
        doc.save(salida)
        return salida

    def _word_pdf(self, origen):
        from docx import Document
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors

        doc_w = Document(origen)
        salida = self._nombre_salida(origen, ".pdf")
        doc_pdf = SimpleDocTemplate(salida, pagesize=A4)
        sty = getSampleStyleSheet()
        els = []

        for para in doc_w.paragraphs:
            if para.text.strip():
                s = "Heading1" if para.style.name.startswith("Heading") else "Normal"
                els.append(Paragraph(para.text, sty[s]))

        for tab in doc_w.tables:
            datos = [[c.text for c in r.cells] for r in tab.rows]
            if datos:
                nc = len(datos[0])
                aw = (A4[0] - 60) / nc
                t = Table(datos, colWidths=[aw]*nc)
                t.setStyle(TableStyle([
                    ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
                    ("FONTSIZE", (0,0), (-1,-1), 8),
                ]))
                els.append(t)

        if not els:
            raise Exception("El archivo Word no tiene contenido extraíble")
        doc_pdf.build(els)
        return salida

    def _word_imgs(self, origen):
        pdf_temp = os.path.join(DOWNLOADS, "_temp_wordpdf.pdf")
        orig_fn = self._nombre_salida
        self._nombre_salida = lambda o, e: pdf_temp
        try:
            self._word_pdf(origen)
        finally:
            self._nombre_salida = orig_fn

        import fitz
        base = Path(origen).stem
        carpeta = os.path.join(DOWNLOADS, f"{base}_imagenes")
        os.makedirs(carpeta, exist_ok=True)
        doc = fitz.open(pdf_temp)
        total = len(doc)
        for i, pag in enumerate(doc, 1):
            pix = pag.get_pixmap(matrix=fitz.Matrix(2, 2))
            pix.save(os.path.join(carpeta, f"pagina_{i:03d}.png"))
        doc.close()
        if os.path.exists(pdf_temp):
            os.remove(pdf_temp)
        return carpeta

    def _word_excel(self, origen):
        import pandas as pd
        from docx import Document
        doc = Document(origen)
        salida = self._nombre_salida(origen, ".xlsx")
        with pd.ExcelWriter(salida, engine="openpyxl") as writer:
            for i, tab in enumerate(doc.tables, 1):
                datos = [[c.text for c in r.cells] for r in tab.rows]
                if datos:
                    df = self._csv_df(";".join(datos[0]) + "\n" +
                                      "\n".join(";".join(r) for r in datos[1:]))
                    if df is not None:
                        df.to_excel(writer, sheet_name=f"Tabla_{i}", index=False)
        return salida


if __name__ == "__main__":
    app = ConvertidorIA()
    app.mainloop()
